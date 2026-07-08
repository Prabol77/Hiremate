from io import BytesIO

from docx import Document
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate


class ExportService:
    """
    Central export service.

    Supports:

    - TXT
    - DOCX
    - PDF
    """

    # =====================================================
    # TXT
    # =====================================================

    @staticmethod
    def export_txt(
        text: str,
    ) -> bytes:
        """
        Export plain text.
        """

        return text.encode("utf-8")

    # =====================================================
    # DOCX
    # =====================================================

    @staticmethod
    def export_docx(
        title: str,
        text: str,
    ) -> bytes:
        """
        Export DOCX.
        """

        document = Document()

        document.add_heading(
            title,
            level=1,
        )

        for paragraph in text.split("\n\n"):

            document.add_paragraph(
                paragraph,
            )

        buffer = BytesIO()

        document.save(buffer)

        buffer.seek(0)

        return buffer.read()

    # =====================================================
    # PDF
    # =====================================================

    @staticmethod
    def export_pdf(
        title: str,
        text: str,
    ) -> bytes:
        """
        Export PDF.
        """

        buffer = BytesIO()

        document = SimpleDocTemplate(
            buffer,
        )

        styles = getSampleStyleSheet()

        story = [

            Paragraph(
                f"<b>{title}</b>",
                styles["Heading1"],
            )
        ]

        for paragraph in text.split("\n\n"):

            story.append(

                Paragraph(
                    paragraph,
                    styles["BodyText"],
                )

            )

        document.build(story)

        buffer.seek(0)

        return buffer.read()